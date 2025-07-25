from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, session, flash
import os
from datetime import datetime
import base64
import wave
import io
import hashlib
import json
from functools import wraps
import re
from dotenv import load_dotenv
import google.generativeai as genai
import speech_recognition as sr
from pydub import AudioSegment
import tempfile
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

# Carregar variáveis de ambiente
load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'sua_chave_secreta_aqui_mude_em_producao')

# Configurar Gemini
gemini_api_key = os.getenv('GEMINI_API_KEY')
if gemini_api_key:
    genai.configure(api_key=gemini_api_key)
    model = genai.GenerativeModel('gemini-2.5-flash')
else:
    model = None
    print("⚠️ AVISO: GEMINI_API_KEY não configurada no arquivo .env")

# Diretórios
RECORDINGS_DIR = 'recordings'
TRANSCRIPTIONS_DIR = 'transcriptions'
USERS_FILE = 'users.json'

if not os.path.exists(RECORDINGS_DIR):
    os.makedirs(RECORDINGS_DIR)
if not os.path.exists(TRANSCRIPTIONS_DIR):
    os.makedirs(TRANSCRIPTIONS_DIR)

# Funções de autenticação
def load_users():
    if os.path.exists(USERS_FILE):
        with open(USERS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_users(users):
    with open(USERS_FILE, 'w', encoding='utf-8') as f:
        json.dump(users, f, ensure_ascii=False, indent=2)

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def sanitize_filename(filename):
    # Remove caracteres especiais e mantém apenas letras, números, espaços e hífens
    return re.sub(r'[^\w\s-]', '', filename).strip()

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# Função para transcrever áudio
def transcribe_audio_with_speech_recognition(audio_path):
    """Transcreve áudio usando uma abordagem mais robusta"""
    try:
        print(f"🔍 Iniciando transcrição de: {audio_path}")
        
        # Verificar se o arquivo existe
        if not os.path.exists(audio_path):
            return "[Erro: Arquivo de áudio não encontrado]"
        
        # Verificar tamanho do arquivo
        file_size = os.path.getsize(audio_path)
        print(f"📊 Tamanho do arquivo: {file_size} bytes")
        
        if file_size < 1000:  # Arquivo muito pequeno
            return "[Erro: Arquivo de áudio muito pequeno ou vazio]"
        
        recognizer = sr.Recognizer()
        
        # Configurações mais robustas para timeout
        recognizer.energy_threshold = 4000
        recognizer.dynamic_energy_threshold = False
        recognizer.pause_threshold = 0.8
        recognizer.operation_timeout = 30  # Aumentado de 10 para 30 segundos
        
        print("📁 Processando arquivo de áudio...")
        
        # Estratégia: Converter para um formato muito básico
        try:
            # Carregar com pydub
            print("🔧 Carregando arquivo com pydub...")
            audio = AudioSegment.from_file(audio_path)
            
            print(f"📊 Propriedades originais: {audio.frame_rate}Hz, {audio.channels} canais, {audio.sample_width*8}bit")
            
            # Verificar duração do áudio
            duration_seconds = len(audio) / 1000.0
            print(f"⏱️ Duração do áudio: {duration_seconds:.1f} segundos")
            
            # Se o áudio for muito longo (mais de 60 segundos), dividir em segmentos
            if duration_seconds > 60:
                print("📂 Áudio longo detectado, dividindo em segmentos...")
                return transcribe_long_audio_in_segments(audio, audio_path)
            
            # Converter para o formato mais básico possível
            # 16kHz, mono, 16-bit (padrão para speech recognition)
            audio = audio.set_frame_rate(16000)
            audio = audio.set_channels(1)
            audio = audio.set_sample_width(2)  # 16-bit
            
            # Normalizar o áudio (ajustar volume)
            audio = audio.normalize()
            
            # Criar arquivo temporário com nome específico
            temp_dir = tempfile.gettempdir()
            temp_path = os.path.join(temp_dir, f"temp_audio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.wav")
            
            print(f"💾 Salvando arquivo temporário: {temp_path}")
            
            # Exportar com configurações específicas
            audio.export(
                temp_path, 
                format="wav",
                parameters=["-acodec", "pcm_s16le"]  # Codec PCM 16-bit little endian
            )
            
            print("🎵 Lendo arquivo convertido...")
            
            # Tentar ler o arquivo convertido
            with sr.AudioFile(temp_path) as source:
                print("🔊 Ajustando para ruído ambiente...")
                recognizer.adjust_for_ambient_noise(source, duration=0.5)
                print("📼 Gravando dados de áudio...")
                audio_data = recognizer.record(source)
            
            print("🗑️ Removendo arquivo temporário...")
            os.unlink(temp_path)
            
        except Exception as conv_error:
            print(f"⚠️ Erro na conversão: {conv_error}")
            
            # Fallback: tentar ler diretamente
            print("🔄 Tentando ler arquivo original...")
            try:
                with sr.AudioFile(audio_path) as source:
                    recognizer.adjust_for_ambient_noise(source, duration=0.5)
                    audio_data = recognizer.record(source)
            except Exception as direct_error:
                print(f"❌ Erro ao ler arquivo: {direct_error}")
                # Último recurso: tentar sem ajuste de ruído
                try:
                    print("🔄 Tentativa final sem ajuste de ruído...")
                    with sr.AudioFile(audio_path) as source:
                        audio_data = recognizer.record(source)
                except Exception as final_error:
                    return f"[Erro: Não foi possível processar o arquivo de áudio. O formato pode não ser suportado. Detalhes: {final_error}]"
        
        # Tentar transcrever com retry automático
        max_retries = 3
        for attempt in range(max_retries):
            try:
                print(f"🤖 Enviando para Google Speech API (tentativa {attempt + 1}/{max_retries})...")
                
                # Usar timeout mais longo na chamada da API
                import socket
                original_timeout = socket.getdefaulttimeout()
                socket.setdefaulttimeout(60)  # 60 segundos para a conexão
                
                text = recognizer.recognize_google(
                    audio_data, 
                    language='pt-BR', 
                    show_all=False
                )
                
                # Restaurar timeout original
                socket.setdefaulttimeout(original_timeout)
                
                print(f"✅ Transcrição concluída: {len(text)} caracteres")
                return text if text.strip() else "[Áudio vazio ou muito baixo]"
                
            except (TimeoutError, socket.timeout) as timeout_error:
                print(f"⏰ Timeout na tentativa {attempt + 1}: {timeout_error}")
                if attempt < max_retries - 1:
                    print(f"🔄 Tentando novamente em 2 segundos...")
                    import time
                    time.sleep(2)
                    continue
                else:
                    return "[Erro: Timeout na transcrição. O arquivo pode ser muito longo ou a conexão está lenta. Tente dividir o áudio em partes menores.]"
                    
            except sr.UnknownValueError:
                print("⚠️ Áudio não compreendido")
                return "[Áudio não pôde ser compreendido - verifique a qualidade do áudio e tente falar mais claramente]"
            except sr.RequestError as e:
                print(f"❌ Erro no serviço Google: {e}")
                return f"[Erro no serviço de reconhecimento: {e}]"
            finally:
                # Garantir que o timeout seja restaurado
                if 'original_timeout' in locals():
                    socket.setdefaulttimeout(original_timeout)
                
    except Exception as e:
        print(f"💥 Erro geral: {str(e)}")
        import traceback
        traceback.print_exc()
        return f"[Erro na transcrição: {str(e)}]"

def transcribe_long_audio_in_segments(audio, original_path):
    """Transcreve áudio longo dividindo em segmentos de 45 segundos"""
    try:
        print("📂 Iniciando transcrição por segmentos...")
        
        segment_length = 45 * 1000  # 45 segundos em millisegundos
        total_duration = len(audio)
        segments = []
        transcriptions = []
        
        # Dividir áudio em segmentos
        for i in range(0, total_duration, segment_length):
            end = min(i + segment_length, total_duration)
            segment = audio[i:end]
            segments.append(segment)
            print(f"📄 Segmento {len(segments)}: {i/1000:.1f}s - {end/1000:.1f}s")
        
        # Transcrever cada segmento
        for idx, segment in enumerate(segments):
            print(f"🎯 Transcrevendo segmento {idx + 1}/{len(segments)}...")
            
            # Salvar segmento temporário
            temp_dir = tempfile.gettempdir()
            segment_path = os.path.join(temp_dir, f"segment_{idx}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.wav")
            
            # Converter segmento para formato adequado
            segment = segment.set_frame_rate(16000).set_channels(1).set_sample_width(2)
            segment = segment.normalize()
            segment.export(segment_path, format="wav", parameters=["-acodec", "pcm_s16le"])
            
            # Transcrever segmento
            recognizer = sr.Recognizer()
            recognizer.operation_timeout = 30
            
            try:
                with sr.AudioFile(segment_path) as source:
                    recognizer.adjust_for_ambient_noise(source, duration=0.3)
                    audio_data = recognizer.record(source)
                
                # Tentar transcrever com timeout
                import socket
                original_timeout = socket.getdefaulttimeout()
                socket.setdefaulttimeout(45)
                
                text = recognizer.recognize_google(audio_data, language='pt-BR')
                socket.setdefaulttimeout(original_timeout)
                
                transcriptions.append(text if text.strip() else "[Segmento silencioso]")
                print(f"✅ Segmento {idx + 1} transcrito: {len(text)} caracteres")
                
            except Exception as seg_error:
                print(f"⚠️ Erro no segmento {idx + 1}: {seg_error}")
                transcriptions.append(f"[Erro no segmento {idx + 1}]")
            finally:
                # Limpar arquivo temporário
                if os.path.exists(segment_path):
                    os.unlink(segment_path)
                if 'original_timeout' in locals():
                    socket.setdefaulttimeout(original_timeout)
        
        # Combinar transcrições
        final_transcription = "\n\n".join([f"[Segmento {i+1}]\n{trans}" for i, trans in enumerate(transcriptions)])
        
        print(f"🎉 Transcrição por segmentos concluída: {len(final_transcription)} caracteres")
        return final_transcription
        
    except Exception as e:
        print(f"💥 Erro na transcrição por segmentos: {e}")
        return f"[Erro na transcrição por segmentos: {str(e)}]"

def improve_transcription_with_gemini(raw_transcription):
    """Melhora a transcrição usando Gemini"""
    if not model or not raw_transcription or raw_transcription.startswith('['):
        return raw_transcription
    
    try:
        prompt = f"""
Você é um assistente especializado em melhorar transcrições médicas. 
Sua tarefa é corrigir e melhorar a seguinte transcrição de uma consulta médica:

Transcrição original:
{raw_transcription}

Por favor:
1. Corrija erros de gramática e ortografia
2. Melhore a pontuação e formatação
3. Organize o texto de forma clara e profissional
4. Mantenha todos os termos médicos e informações importantes
5. Se possível, estruture em seções (ex: Queixa principal, Histórico, Exame físico, etc.)

Retorne apenas o texto melhorado, sem comentários adicionais:
"""
        
        response = model.generate_content(prompt)
        return response.text.strip()
        
    except Exception as e:
        print(f"Erro ao melhorar transcrição com Gemini: {e}")
        return raw_transcription

# Funções para exportação
def create_pdf_from_text(text, title="Resumo da Consulta"):
    """Cria um PDF a partir de texto com formatação melhorada"""
    buffer = io.BytesIO()
    
    # Criar documento PDF
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=72)
    
    # Estilos melhorados
    styles = getSampleStyleSheet()
    
    # Estilo do título principal
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1,  # Centralizado
        textColor=colors.HexColor('#1565c0'),
        fontName='Helvetica-Bold'
    )
    
    # Estilo para cabeçalhos de seção
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=20,
        textColor=colors.HexColor('#2196f3'),
        fontName='Helvetica-Bold'
    )
    
    # Estilo para subcabeçalhos
    subheading_style = ParagraphStyle(
        'CustomSubHeading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=8,
        spaceBefore=15,
        textColor=colors.HexColor('#42a5f5'),
        fontName='Helvetica-Bold'
    )
    
    # Estilo normal melhorado
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        leading=16,
        fontName='Helvetica'
    )
    
    # Estilo para texto em negrito
    bold_style = ParagraphStyle(
        'CustomBold',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        leading=16,
        fontName='Helvetica-Bold'
    )
    
    # Conteúdo
    story = []
    
    # Título com ícone
    story.append(Paragraph(f"📋 {title}", title_style))
    story.append(Spacer(1, 12))
    
    # Data de geração
    story.append(Paragraph(f"📅 Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}", normal_style))
    story.append(Spacer(1, 20))
    
    # Processar o texto linha por linha
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue
            
        # Cabeçalho principal (##)
        if line.startswith('## '):
            clean_line = line[3:].strip()
            story.append(Paragraph(clean_line, heading_style))
        
        # Subcabeçalho (###)
        elif line.startswith('### '):
            clean_line = line[4:].strip()
            # Manter ícones nos subcabeçalhos
            story.append(Paragraph(clean_line, subheading_style))
        
        # Texto em negrito (**texto**)
        elif line.startswith('**') and line.endswith('**'):
            clean_line = line[2:-2].strip()
            story.append(Paragraph(clean_line, bold_style))
        
        # Texto normal
        else:
            story.append(Paragraph(line, normal_style))
    
    # Construir PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_docx_from_text(text, title="Resumo da Consulta"):
    """Cria um DOCX a partir de texto"""
    doc = Document()
    
    # Título
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data de geração
    date_para = doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Espaço
    doc.add_paragraph()
    
    # Converter markdown para texto formatado se necessário
    if '##' in text:
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('## '):
                # Cabeçalho nível 2
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                # Cabeçalho nível 3
                doc.add_heading(line[4:], level=2)
            elif line.startswith('**') and line.endswith('**'):
                # Texto em negrito
                para = doc.add_paragraph()
                para.add_run(line[2:-2]).bold = True
            elif line:
                # Parágrafo normal
                doc.add_paragraph(line)
    else:
        # Texto simples
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para)
    
    # Salvar em buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Rotas de autenticação
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        users = load_users()
        
        if username in users and users[username]['password'] == hash_password(password):
            session['user_id'] = username
            session['user_name'] = users[username]['name']
            flash('Login realizado com sucesso!', 'success')
            return redirect(url_for('index'))
        else:
            flash('Usuário ou senha incorretos!', 'error')
    
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        name = request.form['name']
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        
        users = load_users()
        
        # Validações
        if username in users:
            flash('Nome de usuário já existe!', 'error')
        elif password != confirm_password:
            flash('Senhas não coincidem!', 'error')
        elif len(password) < 6:
            flash('Senha deve ter pelo menos 6 caracteres!', 'error')
        else:
            # Criar novo usuário
            users[username] = {
                'name': name,
                'password': hash_password(password),
                'created_at': datetime.now().isoformat()
            }
            save_users(users)
            
            session['user_id'] = username
            session['user_name'] = name
            flash('Conta criada com sucesso!', 'success')
            return redirect(url_for('index'))
    
    return render_template('register.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('Logout realizado com sucesso!', 'success')
    return redirect(url_for('login'))

# Rotas protegidas
@app.route('/')
@login_required
def index():
    return render_template('index.html', user_name=session.get('user_name'))

def convert_to_wav(audio_bytes):
    """Converte qualquer formato de áudio para WAV compatível com SpeechRecognition"""
    try:
        # Criar arquivo temporário para o áudio original
        with tempfile.NamedTemporaryFile(delete=False, suffix='.webm') as temp_input:
            temp_input.write(audio_bytes)
            temp_input_path = temp_input.name
        
        # Carregar áudio com pydub (suporta vários formatos)
        audio = AudioSegment.from_file(temp_input_path)
        
        # Converter para WAV com configurações específicas para SpeechRecognition
        # 16-bit, mono, 16kHz (formato padrão para reconhecimento de voz)
        audio = audio.set_frame_rate(16000).set_channels(1).set_sample_width(2)
        
        # Criar arquivo temporário para o WAV convertido
        with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_output:
            temp_output_path = temp_output.name
        
        # Exportar como WAV
        audio.export(temp_output_path, format="wav")
        
        # Ler o WAV convertido
        with open(temp_output_path, 'rb') as f:
            wav_bytes = f.read()
        
        # Limpar arquivos temporários
        os.unlink(temp_input_path)
        os.unlink(temp_output_path)
        
        return wav_bytes
        
    except Exception as e:
        print(f"Erro na conversão de áudio: {e}")
        # Se falhar, tentar retornar os bytes originais
        return audio_bytes

@app.route('/save_recording', methods=['POST'])
@login_required
def save_recording():
    try:
        data = request.json
        audio_data = data['audio']
        patient_name = data.get('patient_name', '').strip()
        
        # Remove o prefixo 'data:audio/wav;base64,' ou similar
        if ',' in audio_data:
            audio_data = audio_data.split(',')[1]
        
        # Decodifica o base64
        audio_bytes = base64.b64decode(audio_data)
        
        # Converter para WAV compatível
        print("🔄 Convertendo áudio para formato WAV...")
        wav_bytes = convert_to_wav(audio_bytes)
        
        # Gera nome do arquivo
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        user_id = session['user_id']
        
        if patient_name:
            safe_patient_name = sanitize_filename(patient_name)
            filename = f'{user_id}_{safe_patient_name}_{timestamp}.wav'
        else:
            filename = f'{user_id}_conversa_{timestamp}.wav'
            
        filepath = os.path.join(RECORDINGS_DIR, filename)
        
        # Salvar o WAV convertido
        with open(filepath, 'wb') as f:
            f.write(wav_bytes)
        
        print(f"✅ Áudio convertido e salvo: {filename}")
        
        return jsonify({
            'success': True, 
            'message': 'Gravação salva com sucesso!',
            'filename': filename
        })
    
    except Exception as e:
        print(f"❌ Erro ao salvar gravação: {str(e)}")
        return jsonify({
            'success': False, 
            'message': f'Erro ao salvar gravação: {str(e)}'
        }), 500

@app.route('/transcribe', methods=['POST'])
@login_required
def transcribe_recording():
    try:
        data = request.json
        filename = data['filename']
        user_id = session['user_id']
        
        # Verificar se o arquivo pertence ao usuário
        if not filename.startswith(user_id):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        filepath = os.path.join(RECORDINGS_DIR, filename)
        
        if not os.path.exists(filepath):
            return jsonify({'success': False, 'message': 'Arquivo não encontrado'}), 404
        
        print(f"🎯 Iniciando transcrição de: {filename}")
        
        # Transcrever o áudio
        transcription = transcribe_audio_with_speech_recognition(filepath)
        
        # Melhorar com Gemini se disponível
        if model and not transcription.startswith('['):
            print("🤖 Melhorando transcrição com Gemini...")
            transcription = improve_transcription_with_gemini(transcription)
        
        # Salvar transcrição
        base_filename = os.path.splitext(filename)[0]
        transcription_filename = f'{base_filename}_transcricao.txt'
        transcription_path = os.path.join(TRANSCRIPTIONS_DIR, transcription_filename)
        
        with open(transcription_path, 'w', encoding='utf-8') as f:
            f.write(transcription)
        
        print(f"✅ Transcrição salva: {transcription_filename}")
        
        return jsonify({
            'success': True,
            'message': 'Transcrição concluída com sucesso!',
            'transcription': transcription,
            'transcription_file': transcription_filename
        })
    
    except Exception as e:
        print(f"❌ Erro na transcrição: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Erro na transcrição: {str(e)}'
        }), 500

@app.route('/download_transcription/<filename>')
@login_required
def download_transcription(filename):
    user_id = session['user_id']
    
    # Verificar se o arquivo pertence ao usuário
    if not filename.startswith(user_id):
        flash('Acesso negado!', 'error')
        return redirect(url_for('index'))
    
    filepath = os.path.join(TRANSCRIPTIONS_DIR, filename)
    
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    else:
        flash('Arquivo de transcrição não encontrado!', 'error')
        return redirect(url_for('index'))

@app.route('/rename_recording', methods=['POST'])
@login_required
def rename_recording():
    try:
        data = request.json
        old_filename = data['old_filename']
        new_name = data['new_name']
        user_id = session['user_id']
        
        # Verificar se o arquivo pertence ao usuário
        if not old_filename.startswith(user_id):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        # Sanitizar novo nome
        safe_new_name = sanitize_filename(new_name)
        if not safe_new_name:
            return jsonify({'success': False, 'message': 'Nome inválido'}), 400
        
        # Gerar novo nome do arquivo
        timestamp = old_filename.split('_')[-1]  # Manter timestamp original
        new_filename = f'{user_id}_{safe_new_name}_{timestamp}'
        
        old_path = os.path.join(RECORDINGS_DIR, old_filename)
        new_path = os.path.join(RECORDINGS_DIR, new_filename)
        
        if os.path.exists(old_path):
            os.rename(old_path, new_path)
            
            # Renomear transcrição se existir
            old_transcription = os.path.join(TRANSCRIPTIONS_DIR, 
                                           os.path.splitext(old_filename)[0] + '_transcricao.txt')
            new_transcription = os.path.join(TRANSCRIPTIONS_DIR, 
                                           os.path.splitext(new_filename)[0] + '_transcricao.txt')
            
            if os.path.exists(old_transcription):
                os.rename(old_transcription, new_transcription)
            
            return jsonify({
                'success': True,
                'message': 'Arquivo renomeado com sucesso!',
                'new_filename': new_filename
            })
        else:
            return jsonify({'success': False, 'message': 'Arquivo não encontrado'}), 404
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao renomear: {str(e)}'
        }), 500

@app.route('/transcriptions')
@login_required
def transcriptions_page():
    """Página para listar todas as transcrições"""
    return render_template('transcriptions.html', user_name=session.get('user_name', 'Usuário'))

@app.route('/api/transcriptions')
@login_required
def get_all_transcriptions():
    """API para listar todas as transcrições do usuário"""
    try:
        user_id = session['user_id']
        transcriptions = []
        
        # Listar arquivos de transcrição
        for filename in os.listdir(TRANSCRIPTIONS_DIR):
            if filename.startswith(user_id) and filename.endswith('_transcricao.txt'):
                filepath = os.path.join(TRANSCRIPTIONS_DIR, filename)
                
                # Obter informações do arquivo
                file_size = os.path.getsize(filepath)
                modified_time = os.path.getmtime(filepath)
                modified_date = datetime.fromtimestamp(modified_time).strftime('%d/%m/%Y %H:%M')
                
                # Ler uma prévia do conteúdo
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read()
                        preview = content[:200] + '...' if len(content) > 200 else content
                except:
                    content = '[Erro ao ler arquivo]'
                    preview = '[Erro ao ler arquivo]'
                
                # Extrair nome do paciente do filename
                base_name = filename.replace('_transcricao.txt', '')
                parts = base_name.split('_')
                patient_name = 'Conversa' if len(parts) < 3 or parts[1] == 'conversa' else parts[1]
                
                # Verificar se já tem resumo
                summary_filename = base_name + '_resumo.txt'
                summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_filename)
                has_summary = os.path.exists(summary_path)
                
                # Buscar arquivo WAV original
                wav_filename = base_name + '.wav'
                wav_filepath = os.path.join(RECORDINGS_DIR, wav_filename)
                
                # Obter tamanho do WAV original
                wav_file_size = 0
                if os.path.exists(wav_filepath):
                    wav_file_size = os.path.getsize(wav_filepath)
                
                transcriptions.append({
                    'filename': filename,
                    'patient_name': patient_name,
                    'modified_date': modified_date,
                    'modified_timestamp': modified_time,
                    'size': wav_file_size,
                    'file_size': wav_file_size,
                    'transcription_size': file_size,
                    'preview': preview,
                    'content': content,
                    'has_summary': has_summary,
                    'summary_filename': summary_filename if has_summary else None
                })
        
        # Ordenar por timestamp (mais recente primeiro)
        transcriptions.sort(key=lambda x: x['modified_timestamp'], reverse=True)
        
        return jsonify({
            'success': True,
            'transcriptions': transcriptions
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao listar transcrições: {str(e)}'
        }), 500

@app.route('/api/generate_summary', methods=['POST'])
@login_required
def generate_summary():
    """Gerar resumo de uma transcrição usando IA"""
    try:
        data = request.json
        filename = data['filename']
        custom_prompt = data.get('custom_prompt', '').strip()
        user_id = session['user_id']
        
        # Verificar se o arquivo pertence ao usuário
        if not filename.startswith(user_id):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        filepath = os.path.join(TRANSCRIPTIONS_DIR, filename)
        
        if not os.path.exists(filepath):
            return jsonify({'success': False, 'message': 'Arquivo não encontrado'}), 404
        
        # Ler conteúdo da transcrição
        with open(filepath, 'r', encoding='utf-8') as f:
            transcription_content = f.read()
        
        if not model:
            return jsonify({
                'success': False, 
                'message': 'IA não configurada. Configure a GEMINI_API_KEY no arquivo .env'
            }), 500
        
        # Construir prompt baseado na entrada do usuário
        if custom_prompt:
            # Usar prompt personalizado do usuário
            prompt = f"""
Você é um assistente médico especializado em criar resumos de consultas médicas.
Analise a seguinte transcrição seguindo as instruções específicas do usuário:

INSTRUÇÕES DO USUÁRIO:
{custom_prompt}

Transcrição:
{transcription_content}

Por favor, crie um resumo seguindo exatamente as instruções fornecidas pelo usuário acima.
Mantenha o resumo profissional e focado nos aspectos médicos mais importantes.
"""
        else:
            # Usar prompt padrão
            prompt = f"""
Você é um assistente médico especializado em criar resumos de consultas médicas.
Analise a seguinte transcrição e crie um resumo estruturado e profissional:

Transcrição:
{transcription_content}

Por favor, crie um resumo seguindo esta estrutura:

## RESUMO DA CONSULTA

**Data:** [Extrair se mencionada ou indicar como não especificada]
**Paciente:** [Nome se mencionado ou "Não especificado"]

### 🔍 QUEIXA PRINCIPAL
[Motivo principal da consulta]

### 📋 HISTÓRICO
[Histórico relevante mencionado]

### 🩺 EXAME FÍSICO
[Achados do exame físico se mencionados]

### 💊 CONDUTA/TRATAMENTO
[Medicações, orientações ou tratamentos prescritos]

### 📝 OBSERVAÇÕES IMPORTANTES
[Pontos relevantes adicionais]

### 🔄 RETORNO
[Orientações sobre retorno se mencionadas]

Mantenha o resumo conciso, profissional e focado nos aspectos médicos mais importantes.
"""
        
        try:
            response = model.generate_content(prompt)
            summary = response.text
            
            # Salvar resumo
            base_filename = os.path.splitext(filename)[0]
            summary_filename = f'{base_filename}_resumo.txt'
            summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_filename)
            
            with open(summary_path, 'w', encoding='utf-8') as f:
                f.write(summary)
            
            # Salvar cópia completa automaticamente
            try:
                combined_filename = f'{base_filename}_conversa_completa.txt'
                combined_path = os.path.join(TRANSCRIPTIONS_DIR, combined_filename)
                
                combined_content = f"""# CONVERSA COMPLETA - {datetime.now().strftime('%d/%m/%Y às %H:%M')}

## 📄 TRANSCRIÇÃO ORIGINAL

{transcription_content}

{'='*80}

## 📋 RESUMO GERADO PELA IA

{summary}

{'='*80}

Arquivo gerado automaticamente pelo sistema AudioToText
Contém a transcrição original e o resumo gerado pela IA para referência futura.
"""
                
                with open(combined_path, 'w', encoding='utf-8') as f:
                    f.write(combined_content)
                    
            except Exception as copy_error:
                print(f"Erro ao salvar cópia automática: {copy_error}")
            
            return jsonify({
                'success': True,
                'message': 'Resumo gerado com sucesso!' + (' (personalizado)' if custom_prompt else ''),
                'summary': summary,
                'summary_filename': summary_filename
            })
            
        except Exception as ai_error:
            return jsonify({
                'success': False,
                'message': f'Erro ao gerar resumo com IA: {str(ai_error)}'
            }), 500
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao processar solicitação: {str(e)}'
        }), 500

@app.route('/api/view_summary/<filename>')
@login_required
def view_summary(filename):
    """Visualizar resumo existente"""
    try:
        user_id = session['user_id']
        
        # Verificar se o arquivo pertence ao usuário
        if not filename.startswith(user_id):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        filepath = os.path.join(TRANSCRIPTIONS_DIR, filename)
        
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return jsonify({
                'success': True,
                'content': content
            })
        else:
            return jsonify({'success': False, 'message': 'Resumo não encontrado'}), 404
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao carregar resumo: {str(e)}'
        }), 500

# Novas rotas para exportação
@app.route('/export_summary_pdf/<filename>')
@login_required
def export_summary_pdf(filename):
    """Exportar resumo como PDF"""
    try:
        user_id = session['user_id']
        
        # Verificar se o arquivo pertence ao usuário
        if not filename.startswith(user_id):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        filepath = os.path.join(TRANSCRIPTIONS_DIR, filename)
        
        if not os.path.exists(filepath):
            return jsonify({'success': False, 'message': 'Resumo não encontrado'}), 404
        
        # Ler conteúdo do resumo
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extrair nome do paciente para o título
        base_name = filename.replace('_resumo.txt', '')
        parts = base_name.split('_')
        patient_name = 'Conversa' if len(parts) < 3 or parts[1] == 'conversa' else parts[1]
        title = f"Resumo da Consulta - {patient_name}"
        
        # Criar PDF
        pdf_buffer = create_pdf_from_text(content, title)
        
        # Nome do arquivo PDF
        pdf_filename = f"{base_name}_resumo.pdf"
        
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name=pdf_filename,
            mimetype='application/pdf'
        )
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao exportar PDF: {str(e)}'
        }), 500

@app.route('/export_summary_docx/<filename>')
@login_required
def export_summary_docx(filename):
    """Exportar resumo como DOCX"""
    try:
        user_id = session['user_id']
        
        # Verificar se o arquivo pertence ao usuário
        if not filename.startswith(user_id):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        filepath = os.path.join(TRANSCRIPTIONS_DIR, filename)
        
        if not os.path.exists(filepath):
            return jsonify({'success': False, 'message': 'Resumo não encontrado'}), 404
        
        # Ler conteúdo do resumo
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extrair nome do paciente para o título
        base_name = filename.replace('_resumo.txt', '')
        parts = base_name.split('_')
        patient_name = 'Conversa' if len(parts) < 3 or parts[1] == 'conversa' else parts[1]
        title = f"Resumo da Consulta - {patient_name}"
        
        # Criar DOCX
        docx_buffer = create_docx_from_text(content, title)
        
        # Nome do arquivo DOCX
        docx_filename = f"{base_name}_resumo.docx"
        
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=docx_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao exportar DOCX: {str(e)}'
        }), 500

@app.route('/recordings')
@login_required
def get_recordings():
    try:
        user_id = session['user_id']
        recordings = []
        sessions = {}
        
        # Listar arquivos de gravação
        for filename in os.listdir(RECORDINGS_DIR):
            if filename.startswith(user_id) and filename.endswith('.wav'):
                filepath = os.path.join(RECORDINGS_DIR, filename)
                file_size = os.path.getsize(filepath)
                
                # Verificar se existe transcrição
                transcription_file = os.path.splitext(filename)[0] + '_transcricao.txt'
                transcription_path = os.path.join(TRANSCRIPTIONS_DIR, transcription_file)
                has_transcription = os.path.exists(transcription_path)
                
                if '_sessao_' in filename:
                    # É um segmento de sessão
                    parts = filename.split('_')
                    session_id = parts[2]  # sessao_ID
                    
                    if session_id not in sessions:
                        sessions[session_id] = {
                            'id': session_id,
                            'segments': [],
                            'total_size': 0
                        }
                    
                    sessions[session_id]['segments'].append({
                        'filename': filename,
                        'size': file_size,
                        'has_transcription': has_transcription
                    })
                    sessions[session_id]['total_size'] += file_size
                else:
                    # Gravação simples
                    recordings.append({
                        'filename': filename,
                        'size': file_size,
                        'has_transcription': has_transcription,
                        'type': 'simple'
                    })
        
        # Converter sessões para lista
        session_list = list(sessions.values())
        
        return jsonify({
            'success': True,
            'recordings': recordings,
            'sessions': session_list
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao listar gravações: {str(e)}'
        }), 500

@app.route('/delete_recording', methods=['POST'])
@login_required
def delete_recording():
    try:
        data = request.json
        filename = data['filename']
        user_id = session['user_id']
        
        # Verificar se o arquivo pertence ao usuário
        if not filename.startswith(user_id):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        filepath = os.path.join(RECORDINGS_DIR, filename)
        
        if os.path.exists(filepath):
            os.remove(filepath)
            
            # Remover transcrição se existir
            transcription_file = os.path.splitext(filename)[0] + '_transcricao.txt'
            transcription_path = os.path.join(TRANSCRIPTIONS_DIR, transcription_file)
            
            if os.path.exists(transcription_path):
                os.remove(transcription_path)
            
            # Remover resumo se existir
            summary_file = os.path.splitext(filename)[0] + '_resumo.txt'
            summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_file)
            
            if os.path.exists(summary_path):
                os.remove(summary_path)
            
            return jsonify({
                'success': True,
                'message': 'Arquivo deletado com sucesso!'
            })
        else:
            return jsonify({'success': False, 'message': 'Arquivo não encontrado'}), 404
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao deletar: {str(e)}'
        }), 500

@app.route('/download/<filename>')
@login_required
def download_file(filename):
    user_id = session['user_id']
    
    # Verificar se o arquivo pertence ao usuário
    if not filename.startswith(user_id):
        flash('Acesso negado!', 'error')
        return redirect(url_for('index'))
    
    filepath = os.path.join(RECORDINGS_DIR, filename)
    
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    else:
        flash('Arquivo não encontrado!', 'error')
        return redirect(url_for('index'))

@app.route('/finalize_session', methods=['POST'])
@login_required
def finalize_session():
    try:
        data = request.get_json()
        patient_name = data.get('patient_name', '').strip()
        
        # Obter a sessão atual do usuário
        session_id = session.get('current_session_id')
        
        if not session_id:
            return jsonify({'success': False, 'message': 'Nenhuma sessão ativa encontrada'})
        
        # Atualizar metadados da sessão com nome do paciente
        sessions_dir = os.path.join('recordings', session['user_id'], 'sessions')
        session_dir = os.path.join(sessions_dir, session_id)
        metadata_file = os.path.join(session_dir, 'metadata.json')
        
        if os.path.exists(metadata_file):
            with open(metadata_file, 'r', encoding='utf-8') as f:
                metadata = json.load(f)
            
            metadata['patient_name'] = patient_name
            metadata['finalized_at'] = datetime.now().isoformat()
            
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, ensure_ascii=False, indent=2)
        
        # Limpar sessão atual
        session.pop('current_session_id', None)
        
        message = f'Sessão finalizada com sucesso'
        if patient_name:
            message += f' para o paciente: {patient_name}'
        
        return jsonify({'success': True, 'message': message})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao finalizar sessão: {str(e)}'})

@app.route('/api/save_summary_copy', methods=['POST'])
@login_required
def save_summary_copy():
    """Salvar uma cópia do resumo junto com a conversa original"""
    try:
        data = request.json
        summary_filename = data['summary_filename']
        user_id = session['user_id']
        
        # Verificar se o arquivo pertence ao usuário
        if not summary_filename.startswith(user_id):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_filename)
        
        if not os.path.exists(summary_path):
            return jsonify({'success': False, 'message': 'Resumo não encontrado'}), 404
        
        # Ler conteúdo do resumo
        with open(summary_path, 'r', encoding='utf-8') as f:
            summary_content = f.read()
        
        # Encontrar a transcrição original
        base_name = summary_filename.replace('_resumo.txt', '')
        transcription_filename = base_name + '_transcricao.txt'
        transcription_path = os.path.join(TRANSCRIPTIONS_DIR, transcription_filename)
        
        if not os.path.exists(transcription_path):
            return jsonify({'success': False, 'message': 'Transcrição original não encontrada'}), 404
        
        # Ler conteúdo da transcrição
        with open(transcription_path, 'r', encoding='utf-8') as f:
            transcription_content = f.read()
        
        # Criar arquivo combinado
        combined_filename = base_name + '_conversa_completa.txt'
        combined_path = os.path.join(TRANSCRIPTIONS_DIR, combined_filename)
        
        combined_content = f"""# CONVERSA COMPLETA - {datetime.now().strftime('%d/%m/%Y às %H:%M')}

## 📄 TRANSCRIÇÃO ORIGINAL

{transcription_content}

{'='*80}

## 📋 RESUMO GERADO PELA IA

{summary_content}

{'='*80}

Arquivo gerado automaticamente pelo sistema AudioToText
Contém a transcrição original e o resumo gerado pela IA para referência futura.
"""
        
        # Salvar arquivo combinado
        with open(combined_path, 'w', encoding='utf-8') as f:
            f.write(combined_content)
        
        return jsonify({
            'success': True,
            'message': 'Cópia da conversa completa salva com sucesso!',
            'filename': combined_filename
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao salvar cópia: {str(e)}'
        }), 500

@app.route('/view_transcription/<filename>')
@login_required
def view_transcription(filename):
    try:
        user_id = session['user_id']
        
        # Verificar se o arquivo pertence ao usuário
        if not filename.startswith(user_id):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        filepath = os.path.join(TRANSCRIPTIONS_DIR, filename)
        
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return jsonify({
                'success': True,
                'content': content
            })
        else:
            return jsonify({'success': False, 'message': 'Arquivo de transcrição não encontrado'}), 404
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao carregar transcrição: {str(e)}'
        }), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)